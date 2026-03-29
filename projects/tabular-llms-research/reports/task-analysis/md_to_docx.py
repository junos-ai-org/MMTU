#!/usr/bin/env python3
"""Convert markdown reports to styled .docx files with dark gray text and proper headings.

Supports two modes:
  1. Individual: python md_to_docx.py file1.md file2.md  (one .docx per .md)
  2. Combined:   python md_to_docx.py --combined out.docx file1.md file2.md ...
     Creates a single .docx with a table of contents page, each report on its own page.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
HEADING_COLOR = RGBColor(0x2B, 0x2B, 0x2B)
CODE_GRAY = RGBColor(0x44, 0x44, 0x44)
CODE_BG = "F5F5F5"


def _strip_theme_color(run):
    """Remove theme color attributes so explicit RGB takes effect in Word/Google Docs."""
    rPr = run._element.get_or_add_rPr()
    color_elem = rPr.find(qn("w:color"))
    if color_elem is not None:
        for attr in ["themeColor", "themeShade", "themeTint"]:
            key = qn(f"w:{attr}")
            if key in color_elem.attrib:
                del color_elem.attrib[key]


def set_run_style(run, color=DARK_GRAY, size=Pt(11), font_name="Calibri", bold=False, italic=False):
    run.font.color.rgb = color
    run.font.size = size
    run.font.name = font_name
    run.bold = bold
    run.italic = italic
    _strip_theme_color(run)


def _strip_heading_style_theme(doc):
    """Strip theme colors from built-in heading styles so our RGB sticks."""
    for level in range(1, 5):
        style_name = f"Heading {level}"
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = "Calibri"
            style.font.color.rgb = HEADING_COLOR
            # Strip theme from the style's XML
            rPr = style.element.get_or_add_rPr()
            color_elem = rPr.find(qn("w:color"))
            if color_elem is not None:
                for attr in ["themeColor", "themeShade", "themeTint"]:
                    key = qn(f"w:{attr}")
                    if key in color_elem.attrib:
                        del color_elem.attrib[key]
            # Set heading sizes
            sizes = {1: Pt(22), 2: Pt(16), 3: Pt(13), 4: Pt(11)}
            style.font.size = sizes.get(level, Pt(11))


def add_styled_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = HEADING_COLOR
        run.font.name = "Calibri"
        _strip_theme_color(run)
    return heading


def add_page_break(doc):
    """Insert a page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)


def add_table(doc, header_row, data_rows):
    """Add a formatted table to the document."""
    ncols = len(header_row)
    table = doc.add_table(rows=1, cols=ncols, style="Table Grid")
    table.autofit = True

    # Header
    for i, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text.strip())
        set_run_style(run, color=HEADING_COLOR, size=Pt(10), bold=True)
        # Light gray background for header
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "E8E8E8",
            qn("w:val"): "clear",
        })
        shading.append(shd)

    # Data rows
    for row_data in data_rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            if i < ncols:
                cell = row.cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                text = cell_text.strip()
                # Handle bold markers
                if text.startswith("**") and text.endswith("**"):
                    run = p.add_run(text.strip("*"))
                    set_run_style(run, size=Pt(10), bold=True)
                else:
                    run = p.add_run(text)
                    set_run_style(run, size=Pt(10))

    return table


def parse_inline(paragraph, text, base_size=Pt(11)):
    """Parse inline markdown (bold, italic, code) and add runs to paragraph."""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`|\*[^*]+\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_style(run, size=base_size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_style(run, color=CODE_GRAY, size=Pt(10), font_name="Consolas")
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            set_run_style(run, size=base_size, italic=True)
        else:
            run = paragraph.add_run(part)
            set_run_style(run, size=base_size)


def render_md_lines(doc, lines):
    """Render markdown lines into the document."""
    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip horizontal rules
        if re.match(r'^-{3,}$', line.strip()):
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            add_styled_heading(doc, text, level)
            i += 1
            continue

        # Code blocks
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            code_text = "\n".join(code_lines)
            run = p.add_run(code_text)
            set_run_style(run, color=CODE_GRAY, size=Pt(9), font_name="Consolas")
            pPr = p._element.get_or_add_pPr()
            shd = pPr.makeelement(qn("w:shd"), {
                qn("w:fill"): CODE_BG,
                qn("w:val"): "clear",
            })
            pPr.append(shd)
            continue

        # Tables
        if "|" in line and i + 1 < len(lines) and re.match(r'^\|[-|: ]+\|$', lines[i + 1].strip()):
            header_cells = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2  # skip header and separator
            data_rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                row_cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                data_rows.append(row_cells)
                i += 1
            add_table(doc, header_cells, data_rows)
            doc.add_paragraph()  # spacing after table
            continue

        # Bullet points
        bullet_match = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = bullet_match.group(2)
            p = doc.add_paragraph(style="List Bullet")
            if indent >= 2:
                available = [s.name for s in doc.styles]
                if "List Bullet 2" in available:
                    p.style = doc.styles["List Bullet 2"]
                p.paragraph_format.left_indent = Inches(0.5)
            parse_inline(p, text)
            i += 1
            continue

        # Numbered list
        numbered_match = re.match(r'^(\d+)\.\s+(.*)', line)
        if numbered_match:
            text = numbered_match.group(2)
            p = doc.add_paragraph(style="List Number")
            parse_inline(p, text)
            i += 1
            continue

        # Empty lines
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        parse_inline(p, line)
        i += 1


def _get_section_title(md_path: Path) -> str:
    """Extract the first H1 from a markdown file, or use the filename."""
    with open(md_path) as f:
        for line in f:
            m = re.match(r'^#\s+(.*)', line)
            if m:
                return m.group(1).strip()
    return md_path.stem.replace("_", " ").title()


def create_combined_docx(md_files: list[Path], output_path: Path):
    """Create a single .docx with all reports, each on its own page, with a TOC."""
    doc = Document()

    # Set default styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK_GRAY

    _strip_heading_style_theme(doc)

    # --- Title page / Table of Contents ---
    title = doc.add_heading("Task Analysis Reports", level=0)
    for run in title.runs:
        run.font.color.rgb = HEADING_COLOR
        run.font.name = "Calibri"
        run.font.size = Pt(28)
        _strip_theme_color(run)

    p = doc.add_paragraph()
    run = p.add_run("Encoder vs Decoder Baseline (Large) — Qwen2.5-14B vs T5Gemma-9B")
    set_run_style(run, size=Pt(13), italic=True)

    doc.add_paragraph()  # spacer

    # TOC heading
    toc_heading = doc.add_heading("Contents", level=2)
    for run in toc_heading.runs:
        run.font.color.rgb = HEADING_COLOR
        _strip_theme_color(run)

    # TOC entries
    section_titles = []
    for idx, md_file in enumerate(md_files):
        title_text = _get_section_title(md_file)
        section_titles.append(title_text)
        p = doc.add_paragraph()
        run = p.add_run(f"{idx + 1}.  {title_text}")
        set_run_style(run, size=Pt(12))

    # --- Render each report on its own page ---
    for idx, md_file in enumerate(md_files):
        add_page_break(doc)

        lines = md_file.read_text().splitlines()

        # Add a section number prefix to the first H1
        first_h1_replaced = False
        processed_lines = []
        for line in lines:
            if not first_h1_replaced and re.match(r'^#\s+', line):
                # Replace H1 with numbered version
                h1_text = re.sub(r'^#\s+', '', line)
                processed_lines.append(f"# {idx + 1}. {h1_text}")
                first_h1_replaced = True
            else:
                processed_lines.append(line)

        render_md_lines(doc, processed_lines)

    doc.save(str(output_path))
    print(f"Wrote combined document: {output_path}")


def md_to_docx(md_path: Path, docx_path: Path):
    """Convert a single markdown file to a styled .docx document."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK_GRAY

    _strip_heading_style_theme(doc)

    lines = md_path.read_text().splitlines()
    render_md_lines(doc, lines)

    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")


def main():
    args = sys.argv[1:]

    # Combined mode: --combined output.docx file1.md file2.md ...
    if args and args[0] == "--combined":
        if len(args) < 3:
            print("Usage: md_to_docx.py --combined output.docx file1.md [file2.md ...]")
            sys.exit(1)
        output_path = Path(args[1])
        md_files = [Path(f) for f in args[2:]]
        create_combined_docx(md_files, output_path)
        return

    # Individual mode
    if args:
        md_files = [Path(f) for f in args]
    else:
        # Default: convert all report .md files + synthesis
        report_dir = Path(__file__).parent
        md_files = sorted(report_dir.glob("*_report.md")) + [report_dir / "synthesis.md"]
        md_files = [f for f in md_files if f.exists()]

    # Always generate individual files
    for md_file in md_files:
        docx_file = md_file.with_suffix(".docx")
        md_to_docx(md_file, docx_file)

    # Also generate combined document if multiple files
    if len(md_files) > 1:
        report_dir = md_files[0].parent
        # Put synthesis first, then individual reports
        ordered = []
        synthesis = report_dir / "synthesis.md"
        if synthesis.exists() and synthesis in md_files:
            ordered.append(synthesis)
        for f in md_files:
            if f != synthesis:
                ordered.append(f)
        combined_path = report_dir / "task_analysis_full.docx"
        create_combined_docx(ordered, combined_path)


if __name__ == "__main__":
    main()
